"""
Confluence Connector
=====================
Fetches all content from a Confluence space:
  - Regular wiki pages (HTML → clean text)
  - PDF attachments
  - Word document attachments (.docx)
  - Excel files (.xlsx) — extracts text from cells
"""

import os
import io
import logging
from typing import List, Dict, Any
from dataclasses import dataclass

import httpx
from dotenv import load_dotenv

load_dotenv()
log = logging.getLogger(__name__)


@dataclass
class ConfluenceDocument:
    """A single piece of content from Confluence."""
    id:          str
    title:       str
    content:     str          # clean plain text
    source_url:  str          # link back to original page
    doc_type:    str          # "page" | "pdf" | "docx" | "xlsx"
    space_key:   str
    last_updated: str


class ConfluenceConnector:
    """
    Connects to Confluence REST API and extracts all content.
    
    Usage:
        connector = ConfluenceConnector()
        documents = await connector.fetch_all_documents()
    """

    def __init__(self):
        self.base_url   = os.getenv("CONFLUENCE_URL", "").rstrip("/")
        self.username   = os.getenv("CONFLUENCE_USERNAME", "")
        self.api_token  = os.getenv("CONFLUENCE_API_TOKEN", "")
        self.space_key  = os.getenv("CONFLUENCE_SPACE_KEY", "")
        self.batch_size = int(os.getenv("SYNC_BATCH_SIZE", "10"))

        if not all([self.base_url, self.username, self.api_token, self.space_key]):
            raise ValueError(
                "Missing Confluence config. Set in .env: "
                "CONFLUENCE_URL, CONFLUENCE_USERNAME, "
                "CONFLUENCE_API_TOKEN, CONFLUENCE_SPACE_KEY"
            )

    @property
    def auth(self):
        """Basic auth tuple for httpx."""
        return (self.username, self.api_token)

    @property
    def api_base(self):
        return f"{self.base_url}/rest/api"

    # ── Main entry point ──────────────────────────────────────────────────────

    async def fetch_all_documents(self) -> List[ConfluenceDocument]:
        """
        Fetches ALL content from the configured Confluence space.
        Returns a flat list of ConfluenceDocument objects.
        """
        documents = []

        async with httpx.AsyncClient(auth=self.auth, timeout=60) as client:
            # 1. Fetch all wiki pages
            log.info(f"[Confluence] Fetching pages from space: {self.space_key}")
            pages = await self._fetch_all_pages(client)
            log.info(f"[Confluence] Found {len(pages)} pages")

            for page in pages:
                # Add the page itself as a document
                doc = await self._page_to_document(client, page)
                if doc and doc.content.strip():
                    documents.append(doc)

                # Fetch all attachments on this page
                attachments = await self._fetch_attachments(client, page["id"])
                for attachment in attachments:
                    att_doc = await self._attachment_to_document(
                        client, attachment, page
                    )
                    if att_doc and att_doc.content.strip():
                        documents.append(att_doc)

        log.info(f"[Confluence] Total documents extracted: {len(documents)}")
        return documents

    # ── Page fetching ─────────────────────────────────────────────────────────

    async def _fetch_all_pages(
        self, client: httpx.AsyncClient
    ) -> List[Dict[str, Any]]:
        """Fetches all pages in the space using pagination."""
        pages   = []
        start   = 0
        limit   = 50   # Confluence max per request

        while True:
            response = await client.get(
                f"{self.api_base}/content",
                params={
                    "spaceKey": self.space_key,
                    "type":     "page",
                    "status":   "current",
                    "start":    start,
                    "limit":    limit,
                    "expand":   "version,history",
                }
            )
            response.raise_for_status()
            data = response.json()

            results = data.get("results", [])
            pages.extend(results)

            # Check if there are more pages
            if len(results) < limit:
                break
            start += limit

        return pages

    async def _page_to_document(
        self, client: httpx.AsyncClient, page: Dict
    ) -> ConfluenceDocument | None:
        """Fetches full page content and converts HTML to plain text."""
        try:
            response = await client.get(
                f"{self.api_base}/content/{page['id']}",
                params={"expand": "body.storage,version"}
            )
            response.raise_for_status()
            data = response.json()

            html_content = data.get("body", {}).get("storage", {}).get("value", "")
            clean_text   = self._html_to_text(html_content)

            return ConfluenceDocument(
                id           = page["id"],
                title        = page["title"],
                content      = clean_text,
                source_url   = f"{self.base_url}/wiki/spaces/{self.space_key}/pages/{page['id']}",
                doc_type     = "page",
                space_key    = self.space_key,
                last_updated = data.get("version", {}).get("when", ""),
            )
        except Exception as e:
            log.warning(f"[Confluence] Failed to fetch page {page['id']}: {e}")
            return None

    # ── Attachment fetching ───────────────────────────────────────────────────

    async def _fetch_attachments(
        self, client: httpx.AsyncClient, page_id: str
    ) -> List[Dict]:
        """Gets list of all attachments on a page."""
        try:
            response = await client.get(
                f"{self.api_base}/content/{page_id}/child/attachment",
                params={"limit": 100}
            )
            response.raise_for_status()
            return response.json().get("results", [])
        except Exception as e:
            log.warning(f"[Confluence] Failed to fetch attachments for {page_id}: {e}")
            return []

    async def _attachment_to_document(
        self, client: httpx.AsyncClient,
        attachment: Dict, parent_page: Dict
    ) -> ConfluenceDocument | None:
        """Downloads attachment and extracts text based on file type."""
        filename    = attachment.get("title", "")
        media_type  = attachment.get("metadata", {}).get("mediaType", "")
        download_url = attachment.get("_links", {}).get("download", "")

        if not download_url:
            return None

        # Determine file type
        fname_lower = filename.lower()
        if fname_lower.endswith(".pdf") or "pdf" in media_type:
            doc_type = "pdf"
        elif fname_lower.endswith(".docx") or "wordprocessingml" in media_type:
            doc_type = "docx"
        elif fname_lower.endswith(".xlsx") or "spreadsheetml" in media_type:
            doc_type = "xlsx"
        elif fname_lower.endswith(".txt"):
            doc_type = "txt"
        else:
            # Skip images, videos, zips etc.
            return None

        try:
            # Download the file bytes
            full_url = f"{self.base_url}{download_url}"
            response = await client.get(full_url)
            response.raise_for_status()
            file_bytes = response.content

            # Extract text based on type
            text = self._extract_text(file_bytes, doc_type, filename)

            if not text or not text.strip():
                return None

            return ConfluenceDocument(
                id           = attachment["id"],
                title        = f"{parent_page['title']} — {filename}",
                content      = text,
                source_url   = f"{self.base_url}/wiki/spaces/{self.space_key}/pages/{parent_page['id']}",
                doc_type     = doc_type,
                space_key    = self.space_key,
                last_updated = attachment.get("version", {}).get("when", ""),
            )

        except Exception as e:
            log.warning(f"[Confluence] Failed to process attachment {filename}: {e}")
            return None

    # ── Text extraction ───────────────────────────────────────────────────────

    def _extract_text(self, file_bytes: bytes, doc_type: str, filename: str) -> str:
        """Routes to correct extractor based on file type."""
        try:
            if doc_type == "pdf":
                return self._extract_pdf(file_bytes)
            elif doc_type == "docx":
                return self._extract_docx(file_bytes)
            elif doc_type == "xlsx":
                return self._extract_xlsx(file_bytes)
            elif doc_type == "txt":
                return file_bytes.decode("utf-8", errors="ignore")
            return ""
        except Exception as e:
            log.warning(f"[Parser] Failed to extract text from {filename}: {e}")
            return ""

    def _extract_pdf(self, file_bytes: bytes) -> str:
        """
        Extracts text from PDF using PyMuPDF (fitz).
        Handles scanned PDFs with text layer.
        Free and very fast.
        """
        import fitz  # PyMuPDF
        text_parts = []
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for page_num, page in enumerate(doc, 1):
                text = page.get_text("text")
                if text.strip():
                    text_parts.append(f"[Page {page_num}]\n{text}")
        return "\n\n".join(text_parts)

    def _extract_docx(self, file_bytes: bytes) -> str:
        """
        Extracts text from Word documents (.docx).
        Preserves paragraph structure and table content.
        """
        from docx import Document
        doc        = Document(io.BytesIO(file_bytes))
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        return "\n".join(text_parts)

    def _extract_xlsx(self, file_bytes: bytes) -> str:
        """Extracts text from Excel files."""
        import openpyxl
        wb         = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True)
        text_parts = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"[Sheet: {sheet_name}]")
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(
                    str(cell) for cell in row if cell is not None
                )
                if row_text.strip():
                    text_parts.append(row_text)

        return "\n".join(text_parts)

    def _html_to_text(self, html: str) -> str:
        """Converts Confluence HTML storage format to clean plain text."""
        from bs4 import BeautifulSoup
        if not html:
            return ""
        soup = BeautifulSoup(html, "html.parser")

        # Remove script and style tags
        for tag in soup(["script", "style", "ac:structured-macro"]):
            tag.decompose()

        # Add newlines for block elements
        for tag in soup.find_all(["p", "h1", "h2", "h3", "h4", "li", "tr"]):
            tag.append("\n")

        text = soup.get_text(separator=" ")
        # Clean up excessive whitespace
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return "\n".join(lines)
